# web/services/file_service.py
"""
Сервис для работы с файлами
"""

import os
import json
import zipfile
import base64
import py7zr
import mammoth
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Any
from docx import Document

from core.config.db_config import DATA_PATH, FILES_DIR
from web.utils.logging_helper import logging_helper


class FileService:
    """Сервис для работы с файлами"""
    
    def __init__(self):
        self.data_path = DATA_PATH
        self.files_dir = FILES_DIR
        self.download_stats_file = "file_download_stats.json"
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Создает необходимые директории"""
        Path(self.files_dir).mkdir(parents=True, exist_ok=True)
        
        # Создаем папку для статистики если её нет
        stats_dir = os.path.dirname(self.download_stats_file)
        if stats_dir:
            Path(stats_dir).mkdir(parents=True, exist_ok=True)
    
    def get_all_files(self) -> List[Dict[str, Any]]:
        """Получение списка всех файлов"""
        try:
            with open(self.data_path, encoding='utf-8') as f:
                files = json.load(f)
            
            # Добавляем дополнительную информацию о файлах
            for file_info in files:
                file_path = self.get_file_path(file_info['filename'])
                if os.path.exists(file_path):
                    stat = os.stat(file_path)
                    file_info['size'] = stat.st_size
                    file_info['size_formatted'] = self.format_size(stat.st_size)
                    file_info['modified'] = datetime.fromtimestamp(stat.st_mtime)
                    file_info['exists'] = True
                else:
                    file_info['exists'] = False
                
                # Добавляем метаданные из расширения файла
                ext = os.path.splitext(file_info['filename'])[1]
                file_info['type'] = file_info.get('type', self.get_file_type(ext))
                file_info['icon'] = file_info.get('icon', self.get_file_icon(ext))
                
            return files
            
        except FileNotFoundError:
            logging_helper.log_error(f"Файл данных не найден: {self.data_path}")
            return []
        except json.JSONDecodeError as e:
            logging_helper.log_error(f"Ошибка парсинга JSON файла данных: {str(e)}")
            return []
        except Exception as e:
            logging_helper.log_error(f"Ошибка загрузки файлов: {str(e)}")
            return []
    
    def get_file_by_id(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Получение файла по ID"""
        files = self.get_all_files()
        return next((f for f in files if f['id'] == file_id), None)
    
    def get_file_path(self, filename: str) -> str:
        """Получение полного пути к файлу"""
        return os.path.join(self.files_dir, filename)
    
    def get_file_info(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Получение подробной информации о файле"""
        file_info = self.get_file_by_id(file_id)
        if not file_info:
            return None
        
        file_path = self.get_file_path(file_info['filename'])
        
        # Базовая информация
        info = {
            'id': file_id,
            'filename': file_info['filename'],
            'display_name': file_info.get('display_name', file_info['filename']),
            'category': file_info.get('category', 'Разное'),
            'description': file_info.get('description', ''),
            'type': file_info.get('type'),
            'icon': file_info.get('icon'),
            'exists': os.path.exists(file_path)
        }
        
        if info['exists']:
            try:
                stat = os.stat(file_path)
                info.update({
                    'size': stat.st_size,
                    'size_formatted': self.format_size(stat.st_size),
                    'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    'accessed': datetime.fromtimestamp(stat.st_atime).isoformat()
                })
            except OSError as e:
                logging_helper.log_error(f"Ошибка получения статистики файла {file_path}: {str(e)}")
        
        # Добавляем статистику скачиваний
        download_stats = self.get_download_stats()
        info['downloads'] = download_stats.get(file_id, 0)
        
        return info
    
    def search_files(self, query: str = "", category: str = "", file_type: str = "") -> List[Dict[str, Any]]:
        """Поиск файлов по различным критериям"""
        files = self.get_all_files()
        results = []
        
        query_lower = query.lower() if query else ""
        
        for file_info in files:
            # Проверяем соответствие категории
            if category and file_info.get('category', '').lower() != category.lower():
                continue
            
            # Проверяем соответствие типа файла
            if file_type and file_info.get('type', '').lower() != file_type.lower():
                continue
            
            # Проверяем соответствие поисковому запросу
            if query_lower:
                searchable_text = " ".join([
                    file_info.get('filename', ''),
                    file_info.get('display_name', ''),
                    file_info.get('description', ''),
                    file_info.get('category', '')
                ]).lower()
                
                if query_lower not in searchable_text:
                    continue
            
            results.append(file_info)
        
        return results
    
    def get_categories(self) -> List[Dict[str, Any]]:
        """Получение списка категорий с количеством файлов"""
        files = self.get_all_files()
        categories = {}
        
        for file_info in files:
            category = file_info.get('category', 'Разное')
            if category not in categories:
                categories[category] = {
                    'name': category,
                    'count': 0,
                    'total_size': 0
                }
            categories[category]['count'] += 1
            categories[category]['total_size'] += file_info.get('size', 0)
        
        # Форматируем размеры
        for category_data in categories.values():
            category_data['total_size_formatted'] = self.format_size(category_data['total_size'])
        
        return list(categories.values())
    
    def get_file_stats(self) -> Dict[str, Any]:
        """Получение общей статистики файлов"""
        files = self.get_all_files()
        download_stats = self.get_download_stats()
        
        total_files = len(files)
        total_size = sum(file_info.get('size', 0) for file_info in files)
        total_downloads = sum(download_stats.values())
        
        # Статистика по типам файлов
        type_stats = {}
        for file_info in files:
            file_type = file_info.get('type', 'other')
            if file_type not in type_stats:
                type_stats[file_type] = {'count': 0, 'size': 0}
            type_stats[file_type]['count'] += 1
            type_stats[file_type]['size'] += file_info.get('size', 0)
        
        # Топ загружаемых файлов
        top_downloads = []
        for file_info in files:
            downloads = download_stats.get(file_info['id'], 0)
            if downloads > 0:
                top_downloads.append({
                    'id': file_info['id'],
                    'filename': file_info['filename'],
                    'downloads': downloads
                })
        
        top_downloads.sort(key=lambda x: x['downloads'], reverse=True)
        top_downloads = top_downloads[:10]  # Топ 10
        
        return {
            'total_files': total_files,
            'total_size': total_size,
            'total_size_formatted': self.format_size(total_size),
            'total_downloads': total_downloads,
            'existing_files': len([f for f in files if f.get('exists', False)]),
            'missing_files': len([f for f in files if not f.get('exists', True)]),
            'type_stats': type_stats,
            'top_downloads': top_downloads,
            'categories': len(set(f.get('category', 'Разное') for f in files))
        }
    
    def increment_download_count(self, file_id: str):
        """Увеличение счетчика скачиваний файла"""
        try:
            stats = self.get_download_stats()
            stats[file_id] = stats.get(file_id, 0) + 1
            
            with open(self.download_stats_file, 'w', encoding='utf-8') as f:
                json.dump(stats, f, indent=2, ensure_ascii=False)
                
        except Exception as e:
            logging_helper.log_error(f"Ошибка обновления счетчика скачиваний для {file_id}: {str(e)}")
    
    def get_download_stats(self) -> Dict[str, int]:
        """Получение статистики скачиваний"""
        try:
            if os.path.exists(self.download_stats_file):
                with open(self.download_stats_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return {}
        except Exception as e:
            logging_helper.log_error(f"Ошибка загрузки статистики скачиваний: {str(e)}")
            return {}
    
    def get_file_type(self, extension: str) -> str:
        """Определение типа файла по расширению"""
        mapping = {
            '.zip': 'archive',
            '.rar': 'archive',
            '.7z': 'archive',
            '.tar': 'archive',
            '.gz': 'archive',
            '.docx': 'document',
            '.doc': 'document',
            '.pdf': 'document',
            '.xlsx': 'document',
            '.xls': 'document',
            '.pptx': 'document',
            '.ppt': 'document',
            '.txt': 'text',
            '.md': 'text',
            '.rtf': 'text',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.gif': 'image',
            '.svg': 'image',
            '.bmp': 'image',
            '.ico': 'image',
            '.mp3': 'audio',
            '.wav': 'audio',
            '.flac': 'audio',
            '.aac': 'audio',
            '.mp4': 'video',
            '.avi': 'video',
            '.mkv': 'video',
            '.mov': 'video',
            '.py': 'code',
            '.js': 'code',
            '.html': 'code',
            '.css': 'code',
            '.json': 'code',
            '.xml': 'code',
            '.sql': 'code',
            '.exe': 'executable',
            '.msi': 'executable',
            '.deb': 'executable',
            '.rpm': 'executable'
        }
        return mapping.get(extension.lower(), 'other')
    
    def get_file_icon(self, extension: str) -> str:
        """Определение иконки файла по типу"""
        mapping = {
            '.zip': 'file-archive',
            '.rar': 'file-archive',
            '.7z': 'file-archive',
            '.tar': 'file-archive',
            '.gz': 'file-archive',
            '.docx': 'file-word',
            '.doc': 'file-word',
            '.pdf': 'file-pdf',
            '.xlsx': 'file-excel',
            '.xls': 'file-excel',
            '.pptx': 'file-powerpoint',
            '.ppt': 'file-powerpoint',
            '.txt': 'file-alt',
            '.md': 'file-alt',
            '.rtf': 'file-alt',
            '.jpg': 'file-image',
            '.jpeg': 'file-image',
            '.png': 'file-image',
            '.gif': 'file-image',
            '.svg': 'file-image',
            '.bmp': 'file-image',
            '.ico': 'file-image',
            '.mp3': 'file-audio',
            '.wav': 'file-audio',
            '.flac': 'file-audio',
            '.aac': 'file-audio',
            '.mp4': 'file-video',
            '.avi': 'file-video',
            '.mkv': 'file-video',
            '.mov': 'file-video',
            '.py': 'file-code',
            '.js': 'file-code',
            '.html': 'file-code',
            '.css': 'file-code',
            '.json': 'file-code',
            '.xml': 'file-code',
            '.sql': 'file-code',
            '.exe': 'cog',
            '.msi': 'cog',
            '.deb': 'cog',
            '.rpm': 'cog'
        }
        return mapping.get(extension.lower(), 'file')
    
    def format_size(self, size_bytes: int) -> str:
        """Форматирование размера файла"""
        if size_bytes < 0:
            return "0 B"
            
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"
    
    def format_date(self, date_tuple) -> str:
        """Форматирование даты из tuple в строку"""
        year, month, day, hour, minute, second = date_tuple
        return f"{day:02d}.{month:02d}.{year} {hour:02d}:{minute:02d}"
    
    def create_preview(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Создание превью файла"""
        file_info = self.get_file_by_id(file_id)
        if not file_info:
            return None
        
        file_path = self.get_file_path(file_info['filename'])
        if not os.path.exists(file_path):
            return None
        
        file_ext = os.path.splitext(file_info['filename'])[1].lower()
        
        preview_data = {
            'name': file_info['filename'],
            'type': file_info.get('type', self.get_file_type(file_ext)),
            'icon': file_info.get('icon', self.get_file_icon(file_ext)),
            'content': '',
            'preview_type': 'text',
            'file_id': file_id
        }
        
        try:
            if file_ext == '.zip':
                preview_data.update(self._create_zip_preview(file_path))
            elif file_ext == '.7z':
                preview_data.update(self._create_7z_preview(file_path))
            elif file_ext == '.docx':
                preview_data.update(self._create_docx_preview(file_path))
            elif file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.sql']:
                preview_data.update(self._create_text_preview(file_path))
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.svg', '.bmp']:
                preview_data.update(self._create_image_preview(file_path, file_ext))
            elif file_ext == '.pdf':
                preview_data.update(self._create_pdf_preview(file_id))
            else:
                preview_data['content'] = f'Предварительный просмотр для файлов типа {file_ext} не поддерживается.'
                preview_data['preview_type'] = 'unsupported'
                
        except Exception as e:
            logging_helper.log_error(f"Ошибка создания превью для {file_id}: {str(e)}")
            preview_data['content'] = f'Ошибка создания превью: {str(e)}'
            preview_data['preview_type'] = 'error'
        
        return preview_data
    
    def _create_zip_preview(self, file_path: str) -> Dict[str, Any]:
        """Создание превью для ZIP файлов"""
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            file_list = []
            total_size = 0
            for info in zip_ref.infolist():
                file_size = info.file_size
                total_size += file_size
                file_list.append({
                    'name': info.filename,
                    'size': self.format_size(file_size),
                    'date': self.format_date(info.date_time),
                    'is_dir': info.filename.endswith('/')
                })
            
            return {
                'preview_type': 'zip',
                'content': {
                    'files': file_list,
                    'total_files': len(file_list),
                    'total_size': self.format_size(total_size)
                }
            }
    
    def _create_7z_preview(self, file_path: str) -> Dict[str, Any]:
        """Создание превью для 7Z файлов"""
        try:
            with py7zr.SevenZipFile(file_path, mode='r') as z:
                file_list = []
                total_size = 0
                
                file_infos = z.list()
                
                for file_info in file_infos:
                    filename = file_info.filename
                    is_dir = filename.endswith('/')
                    file_size = file_info.uncompressed if hasattr(file_info, 'uncompressed') else 0
                    total_size += file_size
                    
                    date_info = file_info.creationtime if hasattr(file_info, 'creationtime') else None
                    if date_info:
                        date_str = date_info.strftime('%d.%m.%Y %H:%M')
                    else:
                        date_str = 'Неизвестно'
                    
                    file_list.append({
                        'name': filename,
                        'size': self.format_size(file_size),
                        'date': date_str,
                        'is_dir': is_dir
                    })
                
                return {
                    'preview_type': 'zip',
                    'content': {
                        'files': file_list,
                        'total_files': len(file_list),
                        'total_size': self.format_size(total_size)
                    }
                }
        except Exception as e:
            return {
                'content': f'Ошибка чтения 7z файла: {str(e)}',
                'preview_type': 'error'
            }
    
    def _create_docx_preview(self, file_path: str) -> Dict[str, Any]:
        """Создание превью для DOCX файлов"""
        with open(file_path, 'rb') as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value
            
            doc = Document(file_path)
            core_props = doc.core_properties
            
            return {
                'preview_type': 'docx',
                'content': {
                    'text': text[:2000] + ('...' if len(text) > 2000 else ''),
                    'metadata': {
                        'author': core_props.author or 'Неизвестно',
                        'created': core_props.created.strftime('%d.%m.%Y %H:%M:%S') if core_props.created else 'Неизвестно',
                        'modified': core_props.modified.strftime('%d.%m.%Y %H:%M:%S') if core_props.modified else 'Неизвестно',
                        'title': core_props.title or 'Без названия',
                        'pages': len(doc.paragraphs) // 20 + 1,
                        'paragraphs': len(doc.paragraphs)
                    }
                }
            }
    
    def _create_text_preview(self, file_path: str) -> Dict[str, Any]:
        """Создание превью для текстовых файлов"""
        try:
            # Пробуем различные кодировки
            encodings = ['utf-8', 'cp1251', 'latin1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read(5000)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return {
                    'content': 'Не удалось отобразить содержимое этого текстового файла (неподдерживаемая кодировка).',
                    'preview_type': 'error'
                }
            
            return {
                'content': content + ('...' if len(content) == 5000 else ''),
                'preview_type': 'text'
            }
            
        except Exception as e:
            return {
                'content': f'Ошибка чтения файла: {str(e)}',
                'preview_type': 'error'
            }
    
    def _create_image_preview(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        """Создание превью для изображений"""
        try:
            # Ограничиваем размер изображения для предварительного просмотра
            file_size = os.path.getsize(file_path)
            if file_size > 5 * 1024 * 1024:  # 5MB
                return {
                    'preview_type': 'image_too_large',
                    'content': 'Изображение слишком большое для предварительного просмотра'
                }
            
            with open(file_path, 'rb') as img_file:
                img_data = base64.b64encode(img_file.read()).decode('utf-8')
                mime_type = f"image/{file_ext[1:]}"
                if file_ext == '.svg':
                    mime_type = "image/svg+xml"
                
                return {
                    'preview_type': 'image',
                    'content': f'data:{mime_type};base64,{img_data}'
                }
        except Exception as e:
            return {
                'content': f'Ошибка загрузки изображения: {str(e)}',
                'preview_type': 'error'
            }
    
    def _create_pdf_preview(self, file_id: str) -> Dict[str, Any]:
        """Создание превью для PDF файлов"""
        return {
            'preview_type': 'pdf',
            'content': f'/files/raw/{file_id}'
        }


# Глобальный экземпляр сервиса
file_service = FileService()
